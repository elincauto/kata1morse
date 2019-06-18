from docx import Document
from time import strftime, gmtime
from docx.shared import Pt


document = Document()

document.add_heading('Telegram', 0)

hoy= strftime("%d -%b- %Y", gmtime())

fecha=document.add_paragraph(hoy)

fecha.algnment=1

de=document.add_paragraph()
de.add_run("From:").bold=True
de.add_run("Remitente")


para = document.add_paragraph("")
para.add_run("To:").bold=True
para.add_run("Destinatario")

document.add_heading("Mensaje", level=1)

table = document.add_table(rows=2, cols=1)
table.style="LightShading"

style=table.style
font=style.font
font.name="Courier"
font.size=Pt(12)

#table.style.font.name = "Courier"
#table.style.font.size=Pt(12)

celda_morse = table.rows[0].cells[0]
celda_morse.text=".-.-."

celda_plano=table.rows[1].cells[0]
celda_plano.text="Hola"

document.save("demo.docx")




