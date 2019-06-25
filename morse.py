
from docx import Document
from time import strftime, gmtime
from docx.shared import Pt



morse = {
    'A':'·-', 
    'B':'-···', 
    'C': '-·-·', 
    'D': '-··', 
    'E': '·', 
    'F': '··-·',  
    'G': '--·',
    'H': '····', 
    'I': '··',
    'J': '·---',
    'K': '-·-',
    'L': '·-··',
    'M': '--',
    'N': '-·',
    'Ñ': '--·--', 
    'O': '---',
    'P': '·--·',
    'Q': '--·-',
    'R': '·-·', 
    'S': '···',
    'T': '-',
    'U': '··-',
    'V': '···-',
    'W': '·--',
    'X': '-··-',
    'Y': '-·--', 
    'Z': '--··',
    '0': '-----',
    '1': '·----',
    '2': '··---',
    '3': '···--',
    '4': '····-',
    '5': '·····',
    '6': '-····',
    '7': '--···', 
    '8':'---··',
    '9': '----·',
    '.': '·-·-·-',
    ',': '-·-·--',
    '?': '··--··',
    '"': '·-··-·',
    '!': '--··--'
}


reverso={}
for key in morse:
    valor=morse[key]
    reverso[valor]=key

def toMorse(texto):
    cadena= texto.upper()
    resultado = ""
    for letra in cadena:
        if letra in morse:
            resultado+=morse[letra]
            resultado+=" "
        else:
            resultado+=" "

    return resultado

def toPlain(codigo):
   codigo = codigo.split(' ')
   letras = ''
   for caracter in codigo:
       if caracter in reverso:
           letras += reverso[caracter]

       else:
           letras += " "

   return letras.capitalize()

def Telegram(Destinatario, Remitente, Mensaje):
        document = Document()

        document.add_heading('Telegram', 0)

        fechahora=gmtime()
        
        hoy= strftime("%d -%b- %Y", gmtime())

        fecha=document.add_paragraph(hoy)

        fecha.algnment=1

        de=document.add_paragraph()
        de.add_run("From: ").bold=True
        de.add_run(Remitente)


        para = document.add_paragraph("")
        para.add_run("To: ").bold=True
        para.add_run(Destinatario)

        document.add_heading("Mensaje", level=1)

        table = document.add_table(rows=2, cols=1)
        table.style="LightShading"

        style=table.style
        font=style.font
        font.name="Courier"
        font.size=Pt(12)

        #table.style.font.name = "Courier"
        #table.style.font.size=Pt(12)

        celda_morse = table.rows[0].cells[0]
        celda_morse.text= toMorse(Mensaje)

        celda_plano=table.rows[1].cells[0]
        celda_plano.text=Mensaje

        #document.save("demo.docx")
        
        document.save("envios/{}{}.docx".format(Destinatario, strftime("%Y%m%d%H%M%S%z", fechahora)))